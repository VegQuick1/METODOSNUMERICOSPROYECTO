#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para mejorar y actualizar los manuales con más información
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_table_with_data(doc, data, headers):
    """Añade una tabla con datos específicos"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Encabezados
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Datos
    for row_data in data:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
    
    return table

# ==================== ACTUALIZAR MANUAL TÉCNICO ====================

def enhance_technical_manual():
    """Mejora el manual técnico existente"""
    doc = Document("pdfs/Manual_Tecnico_Metodos_Numericos.docx")
    
    # Añadir más secciones
    doc.add_page_break()
    
    # Nueva sección: Tabla de módulos
    title = doc.add_heading("APÉNDICE A: TABLA COMPARATIVA DE MÓDULOS", level=1)
    
    modules_data = [
        ["Módulo", "Líneas", "Función Principal", "Dependencias"],
        ["main.py", "4", "Punto de entrada", "tkinter, game_app"],
        ["game_data.py", "1345", "Almacén de datos", "ninguna"],
        ["game_app.py", "1335", "Lógica principal", "tkinter, game_data"],
        ["methods_engine.py", "225", "Cálculos numéricos", "numpy"],
        ["methods_mapping.py", "171", "Mapeo de métodos", "ninguna"],
        ["numerical_methods_lessons.py", "variable", "Implementaciones", "numpy, matplotlib"],
        ["music_manager.py", "variable", "Gestor de audio", "pygame"],
    ]
    
    add_table_with_data(doc, modules_data[1:], modules_data[0])
    
    # Nueva sección: Problemas por capítulo
    doc.add_page_break()
    doc.add_heading("APÉNDICE B: DISTRIBUCIÓN DE PROBLEMAS", level=1)
    
    problems_data = [
        ["Capítulo", "Fácil", "Intermedio", "Avanzado", "Final", "Total"],
        ["1. Interpolación", "5", "5", "5", "5", "20"],
        ["2. Ecuaciones Lineales", "5", "5", "5", "5", "20"],
        ["3. Ecuaciones No Lineales", "5", "5", "5", "5", "20"],
        ["4. Integración Numérica", "5", "5", "5", "5", "20"],
        ["5. Mínimos Cuadrados", "5", "6", "6", "6", "23"],
        ["6. EDO", "1", "1", "1", "1", "4"],
        ["TOTAL", "31", "32", "32", "32", "127"],
    ]
    
    add_table_with_data(doc, problems_data[1:], problems_data[0])
    
    # Nueva sección: Especificaciones técnicas
    doc.add_page_break()
    doc.add_heading("APÉNDICE C: ESPECIFICACIONES TÉCNICAS", level=1)
    
    specs_data = [
        ["Especificación", "Valor"],
        ["Lenguaje de Programación", "Python 3.11.9"],
        ["Framework GUI", "Tkinter"],
        ["Base de Datos", "JSON (game_progress.json)"],
        ["Total de Problemas", "127"],
        ["Capítulos", "6"],
        ["Métodos Implementados", "20+"],
        ["Colores Definidos", "10+"],
        ["Tiempos Configurados", "3 (25, 30 min + None)"],
        ["Puntos Máximos Posibles", "1270 (127 × 10)"],
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    for i, text in enumerate(specs_data[0]):
        header_cells[i].text = text
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    for spec in specs_data[1:]:
        row = table.add_row()
        row.cells[0].text = spec[0]
        row.cells[1].text = spec[1]
    
    doc.save("pdfs/Manual_Tecnico_Metodos_Numericos.docx")
    print("✓ Manual Técnico actualizado")

# ==================== ACTUALIZAR MANUAL DE USUARIO ====================

def enhance_user_manual():
    """Mejora el manual de usuario existente"""
    doc = Document("pdfs/Manual_Usuario_Metodos_Numericos.docx")
    
    # Añadir apéndice con estructura completa
    doc.add_page_break()
    doc.add_heading("APÉNDICE: ESTRUCTURA COMPLETA DEL SISTEMA", level=1)
    
    doc.add_heading("6 Capítulos Disponibles", level=2)
    
    chapters_info = {
        "1. Interpolación": [
            "Lagrange - Método de interpolación polinomial",
            "Lineal - Aproximación mediante línea recta",
            "Newton Diferencias Divididas - Método de Newton",
            "Newton Hacia Adelante - Para datos equidistantes",
            "Newton Hacia Atrás - Para datos finales"
        ],
        "2. Ecuaciones Lineales": [
            "Montante - Método Montante",
            "Gauss-Jordán - Eliminación de Gauss-Jordán",
            "Eliminación Gaussiana - Método directo",
            "Gauss-Seidel - Método iterativo",
            "Jacobi - Método de iteración"
        ],
        "3. Ecuaciones No Lineales": [
            "Bisección - Método de bisección",
            "Falsa Posición - Regula-Falsi",
            "Newton-Raphson - Método de Newton",
            "Punto Fijo - Iteración de punto fijo",
            "Secante - Método de la secante"
        ],
        "4. Integración Numérica": [
            "Trapezoidal - Regla de los trapecios",
            "Simpson 1/3 - Regla de Simpson",
            "Simpson 3/8 - Regla compuesta"
        ],
        "5. Mínimos Cuadrados": [
            "Regresión Lineal - Ajuste lineal",
            "Regresión Polinomial - Ajuste polinómico"
        ],
        "6. EDO": [
            "Euler - Método de Euler",
            "Runge-Kutta - Método Runge-Kutta",
            "Predictor-Corrector - Método P-C"
        ]
    }
    
    for chapter, methods in chapters_info.items():
        doc.add_heading(chapter, level=3)
        for method in methods:
            p = doc.add_paragraph(method, style='List Bullet')
    
    doc.save("pdfs/Manual_Usuario_Metodos_Numericos.docx")
    print("✓ Manual de Usuario actualizado")

if __name__ == "__main__":
    os.chdir("c:\\Users\\gokuj\\Downloads\\MN\\METODOSNUMERICOSPROYECTO")
    
    print("Mejorando manuales...")
    enhance_technical_manual()
    enhance_user_manual()
    
    print("\n✓✓✓ Manuales mejorados exitosamente ✓✓✓")
    print("\nArchivos generados:")
    print("- pdfs/Manual_Tecnico_Metodos_Numericos.docx")
    print("- pdfs/Manual_Usuario_Metodos_Numericos.docx")
