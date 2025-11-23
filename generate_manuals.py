#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para generar manuales técnico y de usuario en formato Word
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ==================== FUNCIONES AUXILIARES ====================

def add_heading_style(doc, text, level=1):
    """Añade un encabezado con estilo"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_code_block(doc, code_text, language="Python"):
    """Añade un bloque de código con formato"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    
    # Fondo gris claro para el código
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'E7E6E6')
    p._element.get_or_add_pPr().append(shading_elm)
    
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_table_header(doc, headers):
    """Añade una tabla con encabezados"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Formato de encabezado
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph_format = paragraph.paragraph_format
        
        # Color de fondo azul oscuro
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '1F3A93')
        header_cells[i]._element.get_or_add_pPr().append(shading_elm)
    
    return table

def add_bullet_points(doc, points):
    """Añade puntos de viñeta"""
    for point in points:
        p = doc.add_paragraph(point, style='List Bullet')

def add_bibliography():
    """Retorna la bibliografía del proyecto"""
    return """1. Libro de Texto
   Zamora Pequeño, O., Zamora Pequeño, R. S., & Del Ángel Ramírez, A. 
   (2020). Métodos numéricos aplicados con software (2.ª ed.). Universidad 
   Autónoma de Nuevo León.

2. Documentación de Software (Python)
   Python Software Foundation. (s.f.). tkinter — Python interface to Tcl/Tk. 
   Python 3.12 Documentation. Recuperado el 21 de noviembre de 2025, de 
   https://docs.python.org/3/library/tkinter.html
   
   Python Software Foundation. (s.f.). json — JSON encoder and decoder. 
   Python 3.12 Documentation. Recuperado el 21 de noviembre de 2025, de 
   https://docs.python.org/3/library/json.html
   
   The NumPy Developers. (s.f.). NumPy documentation. NumPy. Recuperado 
   el 21 de noviembre de 2025, de https://numpy.org/doc/
"""

def add_conclusions():
    """Retorna las conclusiones generales del proyecto"""
    return """CONCLUSIONES GENERALES

1. LOGROS ALCANZADOS
   ✓ Sistema educativo interactivo con 127 problemas distribuidos en 6 capítulos
   ✓ Implementación exitosa de métodos numéricos con precisión matemática
   ✓ Interfaz gráfica profesional e intuitiva para usuarios
   ✓ Sistema de progreso y puntuaciones para seguimiento del aprendizaje
   ✓ Base de datos estructurada con problemas validados contra teoría

2. CARACTERÍSTICAS TÉCNICAS DESTACADAS
   • Arquitectura modular y mantenible (separación de datos y lógica)
   • Sistema de tiempo límite adaptativo (Fácil: sin timer, Intermedio: 25 min, 
     Avanzado: 30 min, Final: 25 min)
   • Manejo robusto de errores y excepciones
   • Código bien documentado y estructurado
   • Sistema de respuestas validadas contra documentación de referencia

3. IMPACTO EDUCATIVO
   • Mejora en la comprensión de métodos numéricos mediante práctica interactiva
   • Facilita el aprendizaje progresivo (Fácil → Intermedio → Avanzado)
   • Proporciona retroalimentación inmediata al usuario
   • Permite autoevaluación a través de la Prueba Final

4. POSIBLES MEJORAS FUTURAS
   • Implementación de visualización gráfica en tiempo real
   • Análisis estadístico detallado del desempeño del usuario
   • Integración con sistemas de aprendizaje en línea (LMS)
   • Generación automática de reportes de progreso
   • Versión web usando Flask/Django
   • Soporte multidioma

5. RECOMENDACIONES
   • Usar como herramienta complementaria en cursos de Métodos Numéricos
   • Realizar sesiones de feedback regular con usuarios
   • Actualizar problemas periódicamente con casos reales
   • Considerar gamificación avanzada (insignias, logros, competencias)
"""

# ==================== MANUAL TÉCNICO ====================

def generate_technical_manual():
    """Genera el manual técnico"""
    doc = Document()
    
    # Portada
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MANUAL TÉCNICO\n")
    run.font.size = Pt(28)
    run.font.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Sistema de Aprendizaje Interactivo de Métodos Numéricos\n")
    run.font.size = Pt(16)
    run.italic = True
    
    project_name = doc.add_paragraph()
    project_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project_name.add_run("Métodos Numéricos - Proyecto Educativo")
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Índice
    add_heading_style(doc, "ÍNDICE", level=1)
    toc_items = [
        "1. Introducción",
        "2. Arquitectura del Sistema",
        "3. Módulos Principales",
        "4. Estructura de Datos",
        "5. Código Fuente Destacado",
        "6. Metodología de Desarrollo",
        "7. Validación y Testing",
        "8. Bibliografía",
        "9. Conclusiones"
    ]
    add_bullet_points(doc, toc_items)
    
    doc.add_page_break()
    
    # 1. Introducción
    add_heading_style(doc, "1. INTRODUCCIÓN", level=1)
    doc.add_paragraph(
        "Este documento describe la arquitectura técnica y la implementación del sistema "
        "educativo interactivo para Métodos Numéricos. El proyecto consiste en una "
        "aplicación desarrollada en Python con interfaz gráfica Tkinter que implementa "
        "127 problemas distribuidos en 6 capítulos principales, cubriendo temas esenciales "
        "de cálculo numérico desde el nivel básico hasta avanzado."
    )
    
    add_heading_style(doc, "Objetivos Técnicos", level=2)
    objectives = [
        "Implementar una arquitectura modular y escalable",
        "Garantizar precisión en cálculos matemáticos",
        "Proporcionar una interfaz usuario intuitiva y responsive",
        "Mantener persistencia de datos y progreso del usuario",
        "Validar respuestas contra referencias bibliográficas"
    ]
    add_bullet_points(doc, objectives)
    
    doc.add_page_break()
    
    # 2. Arquitectura del Sistema
    add_heading_style(doc, "2. ARQUITECTURA DEL SISTEMA", level=1)
    doc.add_paragraph(
        "El sistema está diseñado bajo un patrón de separación de responsabilidades, "
        "dividido en tres componentes principales:"
    )
    
    add_heading_style(doc, "Componentes Principales", level=2)
    components = [
        "game_data.py: Capa de datos (127 problemas, estructura de juego)",
        "game_app.py: Lógica de interfaz y control (1335 líneas)",
        "methods_engine.py: Motor de cálculos numéricos (225 líneas)",
        "methods_mapping.py: Mapeo de métodos a funciones (171 líneas)",
        "numerical_methods_lessons.py: Implementaciones educativas de métodos",
        "music_manager.py: Gestor de música y sonidos",
        "additional_methods.py: Métodos numéricos adicionales"
    ]
    add_bullet_points(doc, components)
    
    # Diagrama de flujo
    add_heading_style(doc, "Flujo de Ejecución", level=2)
    flow_steps = [
        "Usuario inicia main.py → Crea instancia de NumericalMethodsGame",
        "Interfaz carga GAME_STRUCTURE desde game_data.py",
        "Usuario selecciona capítulo, método y dificultad",
        "Sistema carga problema correspondiente desde PROBLEM_DATA",
        "Usuario responde → game_app.py valida respuesta",
        "Sistema guarda progreso en game_progress.json",
        "Resultados se muestran y se actualiza puntuación"
    ]
    add_bullet_points(doc, flow_steps)
    
    doc.add_page_break()
    
    # 3. Módulos Principales
    add_heading_style(doc, "3. MÓDULOS PRINCIPALES", level=1)
    
    add_heading_style(doc, "3.1 main.py", level=2)
    doc.add_paragraph(
        "Punto de entrada de la aplicación. Inicializa la ventana principal de Tkinter "
        "y crea una instancia del juego."
    )
    add_code_block(doc, """import tkinter as tk
from game_app import NumericalMethodsGame

if __name__ == "__main__":
    root = tk.Tk()
    app = NumericalMethodsGame(root)
    root.mainloop()
""")
    
    add_heading_style(doc, "3.2 game_data.py (1345 líneas)", level=2)
    doc.add_paragraph(
        "Contiene la definición de todos los problemas, estructura del juego y "
        "datos especiales para el método de Lagrange."
    )
    
    data_sections = [
        "LAGRANGE_FAKE_ANSWERS: Diccionario de opciones falsas para Lagrange",
        "GAME_STRUCTURE: Organización de 6 capítulos con métodos y dificultades",
        "PROBLEM_DATA: 127 problemas con sus preguntas, opciones y respuestas correctas"
    ]
    add_bullet_points(doc, data_sections)
    
    doc.add_paragraph("Estructura de un problema:")
    add_code_block(doc, """{
    'problema_id': {
        'title': 'Enunciado del problema',
        'options': ['opción1', 'opción2', 'opción3', 'opción4'],
        'correct': 'opción_correcta',
        'time_minutes': 25,  # Opcional, None para Fácil
        'table': [...],      # Opcional, datos tabulares
        'x_value': 1.5       # Opcional, valor específico
    }
}
""")
    
    add_heading_style(doc, "3.3 game_app.py (1335 líneas)", level=2)
    doc.add_paragraph(
        "Implementa la lógica principal del juego, interfaz de usuario y control de flujo."
    )
    
    app_features = [
        "Clase NumericalMethodsGame: Controlador principal de la aplicación",
        "Sistema de pantallas: Menú → Capítulos → Métodos → Dificultades → Problemas",
        "Temporizador adaptativo: Fácil (sin timer), Intermedio (25 min), Avanzado (30 min)",
        "Validación de respuestas con retroalimentación inmediata",
        "Persistencia de progreso en JSON",
        "Integración de música y efectos de sonido",
        "Sistema de puntuación y métricas"
    ]
    add_bullet_points(doc, app_features)
    
    doc.add_page_break()
    
    add_heading_style(doc, "3.4 methods_engine.py (225 líneas)", level=2)
    doc.add_paragraph(
        "Motor matemático que implementa los algoritmos numéricos principales."
    )
    
    engine_functions = [
        "solve_lagrange(points, x_to_find): Interpolación de Lagrange",
        "solve_linear_interpolation(a, fa, b, fb, x): Interpolación lineal",
        "solve_newton_divided_differences(points, x_to_find): Diferencias divididas",
        "solve_newton_forward/backward(points, x_to_find): Newton progresivo/regresivo",
        "Integraciones numéricas: Simpson, Trapezoidal, etc."
    ]
    add_bullet_points(doc, engine_functions)
    
    add_code_block(doc, """def solve_lagrange(points, x_to_find):
    \"\"\"Calcula interpolación de Lagrange\"\"\"
    n = len(points)
    g_x = 0
    for i in range(n):
        xi, yi = points[i]
        Li = 1
        for j in range(n):
            if i == j:
                continue
            xj, yj = points[j]
            Li *= (x_to_find - xj) / (xi - xj)
        g_x += yi * Li
    return g_x
""")
    
    add_heading_style(doc, "3.5 methods_mapping.py (171 líneas)", level=2)
    doc.add_paragraph(
        "Mapea métodos matemáticos a sus funciones de visualización y colores."
    )
    
    add_code_block(doc, """METHODS_MAPPING = {
    "Capítulo 1: Interpolación": {
        "Nivel 1: Lagrange": {
            "Intermedio": {
                "function": "_show_lagrange_intermedio",
                "banner_color": "#f8cf39"
            },
            ...
        },
        ...
    },
    ...
}
""")
    
    doc.add_page_break()
    
    # 4. Estructura de Datos
    add_heading_style(doc, "4. ESTRUCTURA DE DATOS", level=1)
    
    add_heading_style(doc, "GAME_STRUCTURE", level=2)
    doc.add_paragraph("Define los 6 capítulos del sistema:")
    
    chapters = [
        "Capítulo 1: Interpolación (5 métodos)",
        "Capítulo 2: Ecuaciones Lineales (5 métodos)",
        "Capítulo 3: Ecuaciones No Lineales (5 métodos)",
        "Capítulo 4: Integración Numérica (3 métodos)",
        "Capítulo 5: Mínimos Cuadrados (2 métodos)",
        "Capítulo 6: EDO (3 métodos)"
    ]
    add_bullet_points(doc, chapters)
    
    add_heading_style(doc, "PROBLEM_DATA", level=2)
    problem_structure = [
        "31 problemas Fácil (teóricos, sin timer)",
        "32 problemas Intermedio (prácticos, 25 minutos)",
        "32 problemas Avanzado (aplicados, 30 minutos)",
        "32 problemas Prueba Final (examen, 25 minutos)",
        "Total: 127 problemas validados"
    ]
    add_bullet_points(doc, problem_structure)
    
    add_heading_style(doc, "LAGRANGE_FAKE_ANSWERS", level=2)
    doc.add_paragraph(
        "Diccionario especial para problemas de Lagrange que requieren seleccionar "
        "componentes de la fórmula:"
    )
    
    add_code_block(doc, """LAGRANGE_FAKE_ANSWERS = {
    'yi': ['y', 'yi+1', 'y(i)', 'yn'],
    'n': ['m', 'k', 'l', 'i'],
    'x-xj': ['x-xi', 'x-x0', 'x-xn', 'xi-x'],
    'xi-xj': ['xi-xk', 'xj-xi', 'xk-xi', 'x-xj'],
    'Σ': ['Σ', 'Π', '∫', 'Δ'],
    'Nada': ['?', '*', '·', '—']
}
""")
    
    doc.add_page_break()
    
    # 5. Código Fuente Destacado
    add_heading_style(doc, "5. CÓDIGO FUENTE DESTACADO", level=1)
    
    add_heading_style(doc, "Lógica del Temporizador (game_app.py)", level=2)
    doc.add_paragraph(
        "Implementación del sistema de tiempo límite adaptativo:"
    )
    add_code_block(doc, """# Obtener tiempo límite del problema
time_min = problem_data.get('time_minutes')

# Crear contenedor de temporizador
timer_container = tk.Frame(frame, bg=COLOR_FONDO)

# No mostrar temporizador para Fácil o sin tiempo
if time_min is None or difficulty.lower() == 'fácil':
    timer_container.pack_forget()
else:
    # Mostrar temporizador para Intermedio/Avanzado
    timer_label = tk.Label(
        timer_container,
        text=f"Tiempo: {time_min}:00",
        font=("Arial", 12, "bold"),
        fg=timer_color,
        bg=COLOR_FONDO
    )
    timer_label.pack()
""")
    
    add_heading_style(doc, "Validación de Respuestas", level=2)
    add_code_block(doc, """def validate_answer(selected, correct):
    \"\"\"Valida si la respuesta del usuario es correcta\"\"\"
    if selected == correct:
        return True, "¡Correcto!"
    else:
        return False, f"Incorrecto. Respuesta correcta: {correct}"
""")
    
    add_heading_style(doc, "Sistema de Persistencia", level=2)
    doc.add_paragraph(
        "El progreso se guarda en game_progress.json en formato JSON:"
    )
    add_code_block(doc, """{
    "chapter": "Interpolación",
    "method": "Lagrange",
    "difficulty": "Intermedio",
    "score": 45,
    "total_attempted": 50,
    "last_played": "2025-11-21 10:30:00"
}
""")
    
    doc.add_page_break()
    
    # 6. Metodología
    add_heading_style(doc, "6. METODOLOGÍA DE DESARROLLO", level=1)
    
    add_heading_style(doc, "Fases de Desarrollo", level=2)
    phases = [
        "Fase 1: Diseño de arquitectura y estructura de datos",
        "Fase 2: Implementación del motor numérico (methods_engine.py)",
        "Fase 3: Desarrollo de interfaz gráfica (game_app.py)",
        "Fase 4: Creación de base de problemas (game_data.py)",
        "Fase 5: Integración y testing",
        "Fase 6: Validación contra documentación de referencia",
        "Fase 7: Optimización e integración de música/sonidos"
    ]
    add_bullet_points(doc, phases)
    
    add_heading_style(doc, "Validación de Respuestas", level=2)
    validation_steps = [
        "Cada respuesta fue validada contra el documento de referencia (Oralia)",
        "Se crearon scripts de verificación (verify_all_times.py, verify_pdf_answers.py)",
        "Pruebas manuales de todos los 127 problemas",
        "Validación de tiempos límite para cada dificultad"
    ]
    add_bullet_points(doc, validation_steps)
    
    doc.add_page_break()
    
    # 7. Testing
    add_heading_style(doc, "7. VALIDACIÓN Y TESTING", level=1)
    
    add_heading_style(doc, "Scripts de Verificación", level=2)
    test_scripts = [
        "verify_all_times.py: Valida que todos los 127 problemas tengan tiempos correctos",
        "verify_pdf_answers.py: Verifica respuestas contra datos del PDF",
        "check_times.py: Auditoría de tiempos por dificultad",
        "apply_simple.py: Batch update de propiedades de problemas"
    ]
    add_bullet_points(doc, test_scripts)
    
    add_heading_style(doc, "Resultados de Testing", level=2)
    results = [
        "✓ 127/127 problemas con sintaxis Python válida",
        "✓ 127/127 problemas con tiempos correctos asignados",
        "✓ 15/15 respuestas PDF validadas y actualizadas",
        "✓ Interfaz responsiva sin errores de compilación",
        "✓ Sistema de persistencia funcionando correctamente"
    ]
    add_bullet_points(doc, results)
    
    doc.add_page_break()
    
    # 8. Bibliografía
    add_heading_style(doc, "8. BIBLIOGRAFÍA", level=1)
    doc.add_paragraph(add_bibliography())
    
    doc.add_page_break()
    
    # 9. Conclusiones
    add_heading_style(doc, "9. CONCLUSIONES", level=1)
    doc.add_paragraph(add_conclusions())
    
    return doc

# ==================== MANUAL DE USUARIO ====================

def generate_user_manual():
    """Genera el manual de usuario"""
    doc = Document()
    
    # Portada
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MANUAL DE USUARIO\n")
    run.font.size = Pt(28)
    run.font.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Sistema de Aprendizaje Interactivo de Métodos Numéricos\n")
    run.font.size = Pt(16)
    run.italic = True
    
    project_name = doc.add_paragraph()
    project_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project_name.add_run("Guía Completa para Estudiantes")
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Índice
    add_heading_style(doc, "ÍNDICE", level=1)
    toc_items = [
        "1. Bienvenida y Descripción General",
        "2. Instalación y Requisitos",
        "3. Cómo Iniciar el Sistema",
        "4. Navegación Básica",
        "5. Estructura de Dificultades",
        "6. Cómo Responder Preguntas",
        "7. Sistema de Puntuación",
        "8. Consejos para el Aprendizaje",
        "9. Solución de Problemas",
        "10. Bibliografía",
        "11. Conclusiones"
    ]
    add_bullet_points(doc, toc_items)
    
    doc.add_page_break()
    
    # 1. Bienvenida
    add_heading_style(doc, "1. BIENVENIDA Y DESCRIPCIÓN GENERAL", level=1)
    doc.add_paragraph(
        "¡Bienvenido al Sistema de Aprendizaje Interactivo de Métodos Numéricos! "
        "Esta aplicación educativa ha sido diseñada especialmente para ayudarte a "
        "aprender y dominar los conceptos fundamentales de los métodos numéricos de forma "
        "interactiva y progresiva."
    )
    
    doc.add_paragraph()
    
    add_heading_style(doc, "¿Qué Aprenderás?", level=2)
    topics = [
        "Métodos de Interpolación (Lagrange, Lineal, Newton)",
        "Resolución de Sistemas de Ecuaciones Lineales",
        "Métodos para Ecuaciones No Lineales",
        "Integración Numérica (Trapezoidal, Simpson)",
        "Aproximación por Mínimos Cuadrados",
        "Resolución de Ecuaciones Diferenciales Ordinarias"
    ]
    add_bullet_points(doc, topics)
    
    add_heading_style(doc, "Contenido del Curso", level=2)
    content = [
        "127 problemas distribuidos en 6 capítulos",
        "4 niveles de dificultad: Fácil, Intermedio, Avanzado, Prueba Final",
        "Validación automática de respuestas",
        "Sistema de progreso y seguimiento",
        "Tiempos límite apropiados para cada nivel"
    ]
    add_bullet_points(doc, content)
    
    doc.add_page_break()
    
    # 2. Instalación
    add_heading_style(doc, "2. INSTALACIÓN Y REQUISITOS", level=1)
    
    add_heading_style(doc, "Requisitos del Sistema", level=2)
    requirements = [
        "Python 3.7 o superior",
        "Sistema Operativo: Windows, macOS o Linux",
        "Memoria RAM: Mínimo 2GB",
        "Espacio en disco: 50MB"
    ]
    add_bullet_points(doc, requirements)
    
    add_heading_style(doc, "Librerías Python Requeridas", level=2)
    libraries = [
        "tkinter (incluido con Python)",
        "numpy: Cálculos matemáticos",
        "PIL (Pillow): Procesamiento de imágenes",
        "pygame: Audio y música"
    ]
    add_bullet_points(doc, libraries)
    
    add_heading_style(doc, "Pasos de Instalación", level=2)
    steps = [
        "1. Descargar el proyecto desde el repositorio",
        "2. Abrir terminal/consola en la carpeta del proyecto",
        "3. Ejecutar: pip install -r requirements.txt",
        "4. Ejecutar: python main.py",
        "5. ¡Disfrutar el aprendizaje!"
    ]
    add_bullet_points(doc, steps)
    
    doc.add_page_break()
    
    # 3. Cómo iniciar
    add_heading_style(doc, "3. CÓMO INICIAR EL SISTEMA", level=1)
    
    add_heading_style(doc, "Inicio Rápido", level=2)
    startup = [
        "Abre una terminal/consola de comandos",
        "Navega hasta la carpeta del proyecto",
        "Ejecuta el comando: python main.py",
        "Se abrirá la ventana principal de la aplicación",
        "¡Estás listo para comenzar!"
    ]
    add_bullet_points(doc, startup)
    
    add_heading_style(doc, "Pantalla Principal", level=2)
    doc.add_paragraph(
        "La pantalla inicial muestra los 6 capítulos disponibles. Cada capítulo "
        "contiene múltiples métodos numéricos para explorar."
    )
    
    main_features = [
        "Título del sistema y bienvenida",
        "6 botones para los capítulos",
        "Barra de progreso general",
        "Indicador de puntuación total",
        "Botón de configuración (opcional)"
    ]
    add_bullet_points(doc, main_features)
    
    doc.add_page_break()
    
    # 4. Navegación
    add_heading_style(doc, "4. NAVEGACIÓN BÁSICA", level=1)
    
    add_heading_style(doc, "Estructura de Menús", level=2)
    doc.add_paragraph(
        "La aplicación organiza el contenido en una estructura jerárquica:"
    )
    
    navigation = [
        "NIVEL 1: Selecciona un Capítulo (Interpolación, Ecuaciones Lineales, etc.)",
        "NIVEL 2: Elige un Método específico dentro del capítulo",
        "NIVEL 3: Selecciona el nivel de dificultad",
        "NIVEL 4: Responde el problema",
        "NIVEL 5: Ve tu resultado y continúa"
    ]
    add_bullet_points(doc, navigation)
    
    add_heading_style(doc, "Botones Principales", level=2)
    buttons_info = [
        "Capítulos: Accede a los 6 temas principales",
        "Métodos: Explora diferentes técnicas dentro de un capítulo",
        "Dificultad: Elige tu nivel (Fácil, Intermedio, Avanzado, Final)",
        "Atrás: Retrocede al menú anterior",
        "Responder: Envía tu respuesta para validación",
        "Siguiente: Continúa con el siguiente problema"
    ]
    add_bullet_points(doc, buttons_info)
    
    doc.add_page_break()
    
    # 5. Dificultades
    add_heading_style(doc, "5. ESTRUCTURA DE DIFICULTADES", level=1)
    
    add_heading_style(doc, "Fácil - Teoría Fundamental", level=2)
    doc.add_paragraph(
        "Los problemas del nivel Fácil son preguntas teóricas que evalúan tu "
        "comprensión de los conceptos fundamentales. NO tienen límite de tiempo."
    )
    easy_features = [
        "Tipo: Preguntas teóricas y conceptuales",
        "Tiempo: Sin límite (trabaja a tu ritmo)",
        "Objetivo: Consolidar conceptos básicos",
        "Ejemplo: ¿Qué requisito tiene el método de Newton-Raphson?",
        "Total: 31 problemas"
    ]
    add_bullet_points(doc, easy_features)
    
    add_heading_style(doc, "Intermedio - Práctica Guiada", level=2)
    doc.add_paragraph(
        "Los problemas del nivel Intermedio son ejercicios prácticos con cálculos. "
        "Dispones de 25 minutos por problema."
    )
    medium_features = [
        "Tipo: Ejercicios prácticos con datos específicos",
        "Tiempo: 25 minutos por problema",
        "Objetivo: Aplicar conceptos en problemas reales",
        "Ejemplo: Calcula la interpolación de Lagrange con los puntos dados",
        "Total: 32 problemas"
    ]
    add_bullet_points(doc, medium_features)
    
    add_heading_style(doc, "Avanzado - Aplicación Profesional", level=2)
    doc.add_paragraph(
        "Los problemas del nivel Avanzado son casos complejos que requieren "
        "integración de múltiples conceptos. Dispones de 30 minutos."
    )
    advanced_features = [
        "Tipo: Problemas integrados y aplicaciones reales",
        "Tiempo: 30 minutos por problema",
        "Objetivo: Dominar técnicas avanzadas",
        "Ejemplo: Encuentra la raíz usando Newton-Raphson en 3 iteraciones",
        "Total: 32 problemas"
    ]
    add_bullet_points(doc, advanced_features)
    
    add_heading_style(doc, "Prueba Final - Evaluación Integral", level=2)
    doc.add_paragraph(
        "La Prueba Final es un examen que integra todos los conceptos aprendidos. "
        "Tienes 25 minutos para completarla."
    )
    final_features = [
        "Tipo: Examen comprensivo de cada capítulo",
        "Tiempo: 25 minutos por examen",
        "Objetivo: Evaluar dominio integral",
        "Requisito: Completar los 3 niveles anteriores",
        "Total: 32 exámenes (uno por método/capítulo)"
    ]
    add_bullet_points(doc, final_features)
    
    doc.add_page_break()
    
    # 6. Cómo responder
    add_heading_style(doc, "6. CÓMO RESPONDER PREGUNTAS", level=1)
    
    add_heading_style(doc, "Tipos de Preguntas", level=2)
    
    add_heading_style(doc, "Opción Múltiple", level=3)
    doc.add_paragraph(
        "La mayoría de problemas son de opción múltiple. Selecciona la respuesta "
        "correcta haciendo clic en el botón correspondiente."
    )
    
    add_heading_style(doc, "Preguntas Especiales (Lagrange)", level=3)
    doc.add_paragraph(
        "Algunos problemas de Lagrange requieren que identifiques componentes de la fórmula. "
        "Lee cuidadosamente qué componente se solicita."
    )
    
    add_heading_style(doc, "Pasos para Responder", level=2)
    answer_steps = [
        "1. Lee el enunciado del problema cuidadosamente",
        "2. Analiza los datos proporcionados (tabla, valores, etc.)",
        "3. Aplica el método numérico apropiado",
        "4. Verifica tu cálculo",
        "5. Selecciona la opción que coincida con tu resultado",
        "6. Haz clic en 'Responder' para validar",
        "7. Revisa la retroalimentación inmediata"
    ]
    add_bullet_points(doc, answer_steps)
    
    add_heading_style(doc, "Consejos para Responder Correctamente", level=2)
    tips = [
        "Lee el problema 2-3 veces antes de responder",
        "Anota los datos importantes en papel",
        "Realiza los cálculos paso a paso",
        "Verifica unidades y precisión numérica",
        "Si dudas, usa el tiempo limite para revisar",
        "Aprende de los errores: revisa la respuesta correcta",
        "Practica múltiples veces para dominar el método"
    ]
    add_bullet_points(doc, tips)
    
    doc.add_page_break()
    
    # 7. Puntuación
    add_heading_style(doc, "7. SISTEMA DE PUNTUACIÓN", level=1)
    
    add_heading_style(doc, "Cómo se Calcula la Puntuación", level=2)
    scoring = [
        "Respuesta Correcta: +10 puntos",
        "Respuesta Incorrecta: 0 puntos",
        "Bonus por Tiempo (si aplica): +5 puntos por problema completado antes de tiempo",
        "Puntuación Total = Suma de todos los problemas completados"
    ]
    add_bullet_points(doc, scoring)
    
    add_heading_style(doc, "Progreso General", level=2)
    progress_info = [
        "Se guarda automáticamente después de cada problema",
        "Puedes ver tu puntuación en la pantalla principal",
        "Tu progreso persiste aunque cierres la aplicación",
        "Puedes revisar problemas anteriores sin perder puntos"
    ]
    add_bullet_points(doc, progress_info)
    
    add_heading_style(doc, "Metas de Aprendizaje", level=2)
    goals = [
        "Fácil: Alcanza 310 puntos (31 × 10)",
        "Intermedio: Alcanza 320 puntos (32 × 10)",
        "Avanzado: Alcanza 320 puntos (32 × 10)",
        "Prueba Final: Alcanza 320 puntos (32 × 10)",
        "Meta Total: 1270 puntos de 127 problemas"
    ]
    add_bullet_points(doc, goals)
    
    doc.add_page_break()
    
    # 8. Consejos
    add_heading_style(doc, "8. CONSEJOS PARA EL APRENDIZAJE", level=1)
    
    add_heading_style(doc, "Estrategia de Estudio Recomendada", level=2)
    strategy = [
        "Semana 1: Completa todos los Fácil (teoría) - 3-4 horas",
        "Semana 2: Completa todos los Intermedio (práctica) - 8-10 horas",
        "Semana 3: Completa todos los Avanzado (aplicación) - 10-12 horas",
        "Semana 4: Completa las Pruebas Finales - 8-10 horas",
        "Total: 30-35 horas de estudio interactivo"
    ]
    add_bullet_points(doc, strategy)
    
    add_heading_style(doc, "Técnicas de Aprendizaje Efectivo", level=2)
    techniques = [
        "Repite cada problema hasta dominarlo completamente",
        "Toma notas de los pasos clave de cada método",
        "Intenta resolver problemas sin ver las opciones",
        "Compara tus cálculos con las respuestas correctas",
        "Identifica patrones en los métodos",
        "Enseña lo aprendido a otros para reforzar",
        "Organiza tu tiempo según el límite disponible"
    ]
    add_bullet_points(doc, techniques)
    
    add_heading_style(doc, "Cuando Tengas Dificultades", level=2)
    difficulties = [
        "Revisa la teoría correspondiente en tus apuntes de clase",
        "Consulta ejemplos resueltos en el material de referencia",
        "Intenta resolver problemas similares paso a paso",
        "Verifica tus cálculos numéricos (precisión decimal)",
        "No dudes en intentar un problema varias veces",
        "Solicita ayuda a tu instructor o compañeros"
    ]
    add_bullet_points(doc, difficulties)
    
    doc.add_page_break()
    
    # 9. Solución de Problemas
    add_heading_style(doc, "9. SOLUCIÓN DE PROBLEMAS COMUNES", level=1)
    
    add_heading_style(doc, "La aplicación no inicia", level=2)
    doc.add_paragraph(
        "Solución: Verifica que Python 3.7+ esté instalado. "
        "Abre terminal en la carpeta del proyecto y ejecuta: python main.py"
    )
    
    add_heading_style(doc, "Módulos no encontrados", level=2)
    doc.add_paragraph(
        "Solución: Instala las dependencias con: pip install -r requirements.txt"
    )
    
    add_heading_style(doc, "Mi respuesta es correcta pero se marca como incorrecta", level=2)
    doc.add_paragraph(
        "Solución: Verifica la precisión decimal de tu respuesta. "
        "Algunos problemas requieren exactitud a ciertos decimales."
    )
    
    add_heading_style(doc, "El temporizador no aparece", level=2)
    doc.add_paragraph(
        "Esto es normal en el nivel Fácil (sin temporizador) y algunos problemas. "
        "El temporizador aparecerá en Intermedio, Avanzado y Prueba Final."
    )
    
    add_heading_style(doc, "Mi progreso no se guarda", level=2)
    doc.add_paragraph(
        "Solución: Asegúrate de que la carpeta tiene permisos de escritura. "
        "Cierra la aplicación correctamente usando el botón Salir."
    )
    
    doc.add_page_break()
    
    # 10. Bibliografía
    add_heading_style(doc, "10. BIBLIOGRAFÍA", level=1)
    doc.add_paragraph(add_bibliography())
    
    doc.add_page_break()
    
    # 11. Conclusiones
    add_heading_style(doc, "11. CONCLUSIONES", level=1)
    doc.add_paragraph(add_conclusions())
    
    return doc

# ==================== GENERACIÓN ====================

if __name__ == "__main__":
    # Rutas absolutas para guardado
    import os
    base_dir = os.path.dirname(os.path.abspath(__file__))
    pdfs_dir = os.path.join(base_dir, "pdfs")
    
    # Crear carpeta pdfs si no existe
    if not os.path.exists(pdfs_dir):
        os.makedirs(pdfs_dir)
    
    print("Generando Manual Técnico...")
    technical_doc = generate_technical_manual()
    technical_path = os.path.join(pdfs_dir, "Manual_Tecnico_Metodos_Numericos.docx")
    technical_doc.save(technical_path)
    print(f"✓ Manual Técnico guardado: {technical_path}")
    
    print("Generando Manual de Usuario...")
    user_doc = generate_user_manual()
    user_path = os.path.join(pdfs_dir, "Manual_Usuario_Metodos_Numericos.docx")
    user_doc.save(user_path)
    print(f"✓ Manual de Usuario guardado: {user_path}")
    
    print("\n✓✓✓ Ambos manuales generados exitosamente ✓✓✓")
